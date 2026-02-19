from flask import Flask, render_template, request, jsonify
from flask_cors import CORS  
import os
import PyPDF2
import docx
import re
import tempfile  
import time
import threading
from datetime import datetime
from werkzeug.utils import secure_filename

app = Flask(__name__)
CORS(app)  

app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# ============================================
# KEEP ALIVE FUNCTION
# ============================================
def keep_alive():
    while True:
        time.sleep(240)
        try:
            print(f"[Keep Alive] Server active at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        except:
            pass

if os.environ.get('RENDER') == 'true':
    thread = threading.Thread(target=keep_alive, daemon=True)
    thread.start()
    print("✅ Keep-alive thread started - Server will not sleep")

# ============================================
# HEALTH CHECK ENDPOINT
# ============================================
@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        'status': 'awake',
        'time': datetime.now().isoformat(),
        'message': 'Server is ready to accept uploads',
        'active_threads': threading.active_count()
    }), 200

# ============================================
# OPTIONS method for CORS preflight
# ============================================
@app.route('/analyze', methods=['OPTIONS'])
def handle_options():
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Accept')
    response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
    response.headers.add('Access-Control-Max-Age', '3600')
    return response

@app.route('/check_ats', methods=['OPTIONS'])
def handle_options_check():
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Accept')
    response.headers.add('Access-Control-Allow-Methods', 'GET, OPTIONS')
    response.headers.add('Access-Control-Max-Age', '3600')
    return response

# ATS Keywords Database
ATS_DATA = {
    'technical_skills': [
        'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
        'typescript', 'go', 'rust', 'scala', 'perl', 'r', 'matlab',
        'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django',
        'flask', 'spring', 'bootstrap', 'jquery', 'ajax', 'rest api', 'graphql',
        'sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'redis', 'elasticsearch',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github',
        'machine learning', 'deep learning', 'ai', 'data science', 'tensorflow',
        'pandas', 'numpy', 'jira', 'excel', 'tableau', 'power bi'
    ],
    
    'soft_skills': [
        'leadership', 'teamwork', 'communication', 'problem solving',
        'critical thinking', 'time management', 'project management',
        'adaptability', 'creativity', 'collaboration', 'analytical',
        'decision making', 'conflict resolution', 'negotiation',
        'presentation', 'public speaking', 'writing', 'interpersonal',
        'emotional intelligence', 'empathy', 'mentoring', 'customer service'
    ],
    
    'action_verbs': [
        'developed', 'managed', 'created', 'implemented', 'designed',
        'led', 'achieved', 'improved', 'increased', 'reduced',
        'built', 'coordinated', 'established', 'generated', 'launched',
        'delivered', 'executed', 'facilitated', 'guided', 'initiated',
        'organized', 'performed', 'planned', 'produced', 'resolved',
        'streamlined', 'strengthened', 'supervised', 'trained', 'transformed',
        'analyzed', 'architected'
    ],
    
    'education_keywords': [
        'bachelor', 'master', 'phd', 'b.tech', 'm.tech', 'b.e.', 'm.e.',
        'b.sc', 'm.sc', 'b.com', 'm.com', 'b.a.', 'm.a.', 'mba', 'bca', 'mca',
        'diploma', 'certification', 'degree', 'university', 'college',
        'gpa', 'cgpa', 'honors', 'distinction', 'merit', 'scholarship'
    ],
    
    'certifications': [
        'aws certified', 'azure certified', 'google certified', 'pmp',
        'prince2', 'scrum master', 'csm', 'psm', 'safe', 'itil',
        'ccna', 'ccnp', 'ceh', 'cissp', 'cisa', 'cism',
        'comptia', 'microsoft certified', 'oracle certified', 'salesforce',
        'tableau certified', 'power bi certified', 'six sigma', 'lean'
    ]
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text

def extract_text_from_docx(filepath):
    text = ""
    try:
        doc = docx.Document(filepath)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text

def extract_achievements(text):
    """Extract achievements from resume text"""
    achievements = {
        'has_achievements': False,
        'achievement_count': 0,
        'achievement_list': [],
        'achievement_score': 0
    }
    
    text_lower = text.lower()
    lines = text.split('\n')
    
    # Achievement-related keywords
    achievement_keywords = [
        'achievement', 'award', 'won', 'secured', 'rank', 'position',
        'certificate', 'recognition', 'honor', 'medal', 'trophy',
        'scholarship', 'distinction', 'merit', 'excellence',
        'first prize', 'second prize', 'third prize', 'winner',
        'champion', 'gold medal', 'silver medal', 'bronze medal',
        'topper', 'outstanding', 'exceptional', 'performance'
    ]
    
    # Look for achievement section
    achievement_section_found = False
    current_section = ""
    
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        
        # Check if this line starts an achievement section
        if any(keyword in line_lower for keyword in ['achievement', 'award', 'honor', 'recognition']):
            if ':' in line or len(line.split()) < 8:  # Likely a section header
                achievement_section_found = True
                current_section = "achievements"
                continue
        
        # If we're in achievement section, collect achievements
        if achievement_section_found and current_section == "achievements":
            if line.strip() and not any(keyword in line_lower for keyword in ['skill', 'education', 'project', 'experience']):
                # Check if this line contains achievement keywords or bullet points
                if any(keyword in line_lower for keyword in achievement_keywords) or line.strip().startswith(('•', '-', '*', '✓')):
                    achievements['achievement_list'].append(line.strip())
    
    # Also scan entire text for achievement patterns
    achievement_patterns = [
        r'secured\s+(\d+\w*\s+)?rank',
        r'won\s+(\d+\w*\s+)?prize',
        r'secured\s+\d+\s*%',
        r'achieved\s+\d+',
        r'ranked\s+#?\d+',
        r'position\s+#?\d+',
        r'topper\s+in',
        r'outstanding\s+performance',
        r'excellence\s+award'
    ]
    
    for pattern in achievement_patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            achievements['achievement_count'] += len(matches)
    
    # Count achievements from list
    achievements['achievement_count'] += len(achievements['achievement_list'])
    
    # Check for bullet points that might contain achievements
    bullet_lines = re.findall(r'[•\-*✓]\s*(.*?)(?:\n|$)', text)
    for line in bullet_lines:
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in achievement_keywords):
            if line not in achievements['achievement_list']:
                achievements['achievement_list'].append(line.strip())
                achievements['achievement_count'] += 1
    
    # Calculate achievement score
    if achievements['achievement_count'] >= 3:
        achievements['achievement_score'] = 10
    elif achievements['achievement_count'] == 2:
        achievements['achievement_score'] = 7
    elif achievements['achievement_count'] == 1:
        achievements['achievement_score'] = 4
    else:
        achievements['achievement_score'] = 0
    
    achievements['has_achievements'] = achievements['achievement_count'] > 0
    
    return achievements

def extract_contact_info(text):
    """Extract contact information from resume - QR codes based on text only"""
    info = {
        'email': [],
        'phone': [],
        'linkedin': False,
        'github': False,
        'linkedin_qr': False,
        'github_qr': False,
        'qr_code': False,
        'portfolio': False
    }
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    info['email'] = list(set(emails))
    
    # Phone patterns (Indian numbers with spaces)
    phone_patterns = [
        r'\b(?:\+91|0)?[6-9]\d{9}\b',
        r'\b(?:\+91|0)?[6-9]\d{4}\s\d{5}\b',
        r'\b(?:\+91|0|91)\s[6-9]\d{9}\b',
        r'\b[6-9]\d{4}-\d{5}\b',
        r'\+91[6-9]\d{9}\b',
    ]
    
    phones = []
    for pattern in phone_patterns:
        found = re.findall(pattern, text)
        phones.extend(found)
    
    cleaned_phones = []
    for phone in phones:
        clean = re.sub(r'[\s\-]', '', phone)
        if clean.startswith('+91'):
            clean = clean[3:]
        elif clean.startswith('91') and len(clean) > 10:
            clean = clean[2:]
        elif clean.startswith('0') and len(clean) > 10:
            clean = clean[1:]
        
        if len(clean) == 10 and clean[0] in '6789':
            cleaned_phones.append(clean)
    
    info['phone'] = list(set(cleaned_phones))
    
    # QR Code Detection
    text_lower = text.lower()
    
    # Check for general QR code mentions
    if re.search(r'\bqr\b|\bqrcode\b|\bqr code\b', text_lower):
        info['qr_code'] = True
    
    # LinkedIn QR Detection
    linkedin_qr_patterns = [
        r'linkedin\s*qr',
        r'linkedin\s*:',
        r'linkedin',
        r'li\s*qr',
        r'in\s*qr',
    ]
    
    for pattern in linkedin_qr_patterns:
        if re.search(pattern, text_lower):
            info['linkedin'] = True
            info['linkedin_qr'] = True
            info['qr_code'] = True
            break
    
    # GitHub QR Detection
    github_qr_patterns = [
        r'github\s*qr',
        r'github\s*:',
        r'github',
        r'gh\s*qr',
    ]
    
    for pattern in github_qr_patterns:
        if re.search(pattern, text_lower):
            info['github'] = True
            info['github_qr'] = True
            info['qr_code'] = True
            break
    
    # Regular URLs
    if re.search(r'linkedin\.com/in/[\w\-]+', text_lower):
        info['linkedin'] = True
    
    if re.search(r'github\.com/[\w\-]+', text_lower):
        info['github'] = True
    
    # Portfolio detection
    portfolio_patterns = [r'portfolio', r'personal website', r'my website']
    for pattern in portfolio_patterns:
        if re.search(pattern, text_lower) and not info['github'] and not info['linkedin']:
            info['portfolio'] = True
            break
    
    return info

def calculate_ats_score(text):
    """Calculate ATS score based on various factors"""
    text_lower = text.lower()
    words = text.split()
    
    # Find keywords
    found_keywords = {
        'technical_skills': [],
        'soft_skills': [],
        'action_verbs': [],
        'education_keywords': [],
        'certifications': []
    }
    
    for category, keywords in ATS_DATA.items():
        for keyword in keywords:
            if keyword.lower() in text_lower:
                found_keywords[category].append(keyword)
    
    # Calculate category scores
    category_scores = {}
    total_possible = 0
    total_found = 0
    
    for category, keywords in ATS_DATA.items():
        possible = len(keywords)
        found = len(found_keywords[category])
        total_possible += possible
        total_found += found
        category_scores[category] = round((found / possible) * 100, 1) if possible > 0 else 0
    
    # Overall ATS score
    ats_score = round((total_found / total_possible) * 100, 1) if total_possible > 0 else 0
    
    # Extract contact information
    contact_info = extract_contact_info(text)
    
    # Extract achievements
    achievements = extract_achievements(text)
    
    # Contact information score
    contact_score = 0
    if contact_info['email']:
        contact_score += 3
    if contact_info['phone']:
        contact_score += 3
    if contact_info['linkedin']:
        contact_score += 2
    if contact_info['github']:
        contact_score += 2
    if contact_info['linkedin_qr']:
        contact_score += 2
    if contact_info['github_qr']:
        contact_score += 2
    if contact_info['qr_code']:
        contact_score += 1
    
    # Length score
    word_count = len(words)
    if word_count < 300:
        length_score = 2
    elif word_count < 400:
        length_score = 5
    elif word_count < 600:
        length_score = 8
    elif word_count < 800:
        length_score = 10
    else:
        length_score = 7
    
    # Final score (including achievements)
    final_score = round((ats_score * 0.7) + (contact_score * 2) + (length_score * 0.5) + (achievements['achievement_score']), 1)
    
    return {
        'final_score': final_score,
        'keyword_score': ats_score,
        'contact_score': contact_score,
        'length_score': length_score,
        'achievements': achievements,
        'category_scores': category_scores,
        'found_keywords': found_keywords,
        'keyword_counts': {
            'total': total_found,
            'technical': len(found_keywords['technical_skills']),
            'soft': len(found_keywords['soft_skills']),
            'actions': len(found_keywords['action_verbs']),
            'education': len(found_keywords['education_keywords']),
            'certifications': len(found_keywords['certifications'])
        },
        'contact_info': contact_info,
        'word_count': word_count
    }

def get_score_rating(score):
    if score >= 90:
        return "Excellent", "🎯 Top 5% Resume"
    elif score >= 80:
        return "Very Good", "📈 Strong Contender"
    elif score >= 70:
        return "Good", "👍 Above Average"
    elif score >= 60:
        return "Average", "📊 Needs Optimization"
    elif score >= 50:
        return "Below Average", "⚠️ Improve Keywords"
    else:
        return "Poor", "❌ Major Changes Needed"

def generate_recommendations(analysis, text):
    recommendations = []
    found = analysis['found_keywords']
    counts = analysis['keyword_counts']
    contact = analysis['contact_info']
    achievements = analysis['achievements']
    
    # Technical skills
    if counts['technical'] < 5:
        recommendations.append({
            'category': 'Technical Skills',
            'priority': 'High',
            'message': 'Add more technical skills relevant to your target role.',
            'examples': ['Python', 'Java', 'SQL', 'AWS', 'Docker', 'React']
        })
    
    # Soft skills
    if counts['soft'] < 3:
        recommendations.append({
            'category': 'Soft Skills',
            'priority': 'Medium',
            'message': 'Include soft skills that employers look for.',
            'examples': ['Leadership', 'Communication', 'Problem Solving']
        })
    
    # Action verbs
    if counts['actions'] < 5:
        recommendations.append({
            'category': 'Action Verbs',
            'priority': 'High',
            'message': 'Use strong action verbs to describe your achievements.',
            'examples': ['Developed', 'Managed', 'Implemented', 'Led', 'Achieved']
        })
    
    # Contact info
    if not contact['email']:
        recommendations.append({
            'category': 'Contact Info',
            'priority': 'Critical',
            'message': 'Add your email address to your resume.',
            'examples': ['your.name@email.com']
        })
    
    if not contact['phone']:
        recommendations.append({
            'category': 'Contact Info',
            'priority': 'High',
            'message': 'Include an Indian mobile number for recruiters.',
            'examples': ['+91 98765 43210', '9876543210']
        })
    
    if not contact['linkedin']:
        recommendations.append({
            'category': 'Professional Links',
            'priority': 'Medium',
            'message': 'Add your LinkedIn profile URL or QR code.',
            'examples': ['LinkedIn QR code with text "LinkedIn QR"']
        })
    
    if not contact['github']:
        recommendations.append({
            'category': 'Developer Links',
            'priority': 'Medium',
            'message': 'Add your GitHub profile URL or QR code.',
            'examples': ['GitHub QR code with text "GitHub QR"']
        })
    
    # Achievement recommendations
    if not achievements['has_achievements']:
        recommendations.append({
            'category': 'Achievements',
            'priority': 'Medium',
            'message': 'Add your achievements to stand out!',
            'examples': ['Secured 1st Rank in competition', 'Won coding challenge', 'Achieved 90%+ in exams']
        })
    elif achievements['achievement_count'] < 3:
        recommendations.append({
            'category': 'Achievements',
            'priority': 'Medium',
            'message': f'You have {achievements["achievement_count"]} achievement(s). Add more to boost your profile!',
            'examples': ['Add awards', 'Add recognitions', 'Add competition wins']
        })
    
    return recommendations[:7]

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST', 'OPTIONS'])
def analyze():
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Accept')
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        return response

    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['resume']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Please upload PDF or DOCX file'}), 400
    
    try:
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(filepath)
        
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(filepath)
        else:
            text = extract_text_from_docx(filepath)
        
        os.remove(filepath)
        
        if not text or len(text.strip()) < 50:
            return jsonify({'error': 'Could not extract enough text from file.'}), 400
        
        analysis = calculate_ats_score(text)
        rating, description = get_score_rating(analysis['final_score'])
        analysis['rating'] = rating
        analysis['rating_description'] = description
        analysis['recommendations'] = generate_recommendations(analysis, text)
        analysis['sample_keywords'] = {
            'technical_skills': ATS_DATA['technical_skills'][:10],
            'soft_skills': ATS_DATA['soft_skills'][:8],
            'action_verbs': ATS_DATA['action_verbs'][:10]
        }
        
        response = jsonify(analysis)
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
    
    except Exception as e:
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/check_ats', methods=['GET', 'OPTIONS'])
def check_ats():
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Accept')
        response.headers.add('Access-Control-Allow-Methods', 'GET, OPTIONS')
        return response

    response = jsonify({
        'message': "🔍 Amazon ATS (Applicant Tracking System) Test",
        'tips': [
            {
                'title': 'Formatting Tips',
                'items': [
                    'Use standard fonts (Arial, Calibri, Times New Roman)',
                    'Save as PDF or DOCX (not JPEG or PNG)',
                    'Avoid headers, footers, tables, and columns',
                    'Use bullet points for easy scanning'
                ]
            },
            {
                'title': 'Content Tips',
                'items': [
                    'Include job-specific keywords from description',
                    'Quantify achievements with numbers and percentages',
                    'Use both technical skills and soft skills',
                    'Add complete contact information with Indian mobile number'
                ]
            },
            {
                'title': 'Achievement Tips',
                'items': [
                    'Add your achievements in a dedicated section',
                    'Include awards, ranks, and competition wins',
                    'Quantify your achievements with numbers',
                    'Example: "Secured 1st Rank in typing competition"'
                ]
            },
            {
                'title': 'QR Code Tips',
                'items': [
                    'Add QR codes for LinkedIn and GitHub',
                    'Write simple text like "LinkedIn QR" and "GitHub QR" near the codes',
                    'The system will detect them automatically',
                    'QR codes make your resume stand out!'
                ]
            }
        ]
    })
    
    response.headers.add('Access-Control-Allow-Origin', '*')
    return response

if __name__ == '__main__':
    print("🚀 ATS Resume Analyzer Starting...")
    print("📱 Indian Mobile Number Support Enabled")
    print("📲 QR Code Detection Enabled (Text-based)")
    print("🏆 Achievement Scanning Enabled")
    print("🐙 GitHub QR Detection with simple text")
    print("🔗 LinkedIn QR Detection with simple text")
    print(f"📁 Upload folder: {app.config['UPLOAD_FOLDER']}")
    print("✨ Ready to analyze resumes!")
    
    app.run(host='0.0.0.0', port=5000, debug=True)